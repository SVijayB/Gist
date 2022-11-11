from docx import Document
from docx2pdf import convert
from datetime import datetime
import os
import re


def result_generation(data):
    document = Document()

    p = document.add_paragraph()
    p = p.insert_paragraph_before("                    ")
    r = p.add_run()
    r.add_picture("assets/logo.jpeg")

    ID = open("assets/ID.txt", "r").read()
    now = datetime.now()
    DateAndTime = now.strftime("%d/%m/%Y %H:%M:%S")

    updated_ID = int(ID) + 1
    update = open("assets/ID.txt", "w")
    update.write(str(updated_ID))
    update.close()

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = (
        "ID:"
        + str("{num:0>4}".format(num=str(ID)))
        + " Generated at : "
        + str(DateAndTime)
    )
    paragraph.style = document.styles["Header"]

    document.add_heading("Results", 0)

    document.add_heading("File 1 Details")
    filename1 = re.findall(r"([\w\d\-.]+\.(?:txt|pdf|docx))", file1)
    document.add_paragraph("NAME : " + str(filename1[0]))
    document.add_paragraph("Uploaded Date : " + str(curr_time1.strftime("%d/%m/%Y")))
    document.add_paragraph("Uploaded Time : " + str(curr_time1.strftime("%H:%M:%S")))
    file_stats = os.stat(file1)
    file_size = str(f"File Size : {round(file_stats.st_size / (1024 * 1024),2)} MB")
    document.add_paragraph(file_size)

    document.add_heading("File 2 Details")
    filename2 = re.findall(r"([\w\d\-.]+\.(?:txt|pdf|docx))", file2)
    document.add_paragraph("NAME : " + str(filename2[0]))
    now = datetime.now()
    date = now.strftime("%d/%m/%Y")
    time = now.strftime("%H:%M:%S")
    document.add_paragraph("Uploaded Date : " + str(curr_time2.strftime("%d/%m/%Y")))
    document.add_paragraph("Uploaded Time : " + str(curr_time2.strftime("%H:%M:%S")))
    file_stats = os.stat(file2)
    file_size = str(f"File Size : {round(file_stats.st_size / (1024 * 1024),2)} MB")
    document.add_paragraph(file_size)
    document.add_paragraph("\n")

    document.add_heading("Similarity Rates")
    document.add_paragraph("Jaccard : " + str(jaccard))
    document.add_paragraph("Dice : " + str(dice))

    document.add_heading("Execution Time")
    document.add_paragraph("Time taken : " + str(execution_time) + " seconds")

    document.add_page_break()

    document.add_heading("Report", 0)

    doc_location = "temp\Result_{num:0>4}.docx".format(num=str(ID))
    document.save(doc_location)
    convert(doc_location, output_path="temp")
    os.remove(doc_location)


data = {
    "summary": "Hello there",
    "title": "Hello",
    "content": "Nice to meet you",
    "type": "Image",
}
result_generation(data)
