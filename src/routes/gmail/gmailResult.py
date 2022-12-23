from pprint import pprint
import time
from src.components import summarizer
import json
from docx import Document
from docx2pdf import convert
from datetime import datetime
import pythoncom
import os
import re
import win32com.client


def GmailSummarizer(responses, email_address):
    start = time.time()
    # for email in responses:
    #     email["content"] = email["content"].replace("\n", " ")
    #     email["content"] = email["content"].replace("\r", " ")
    #     email["content"] = email["content"].replace("\t", " ")
    #     summarizer.summarize(email, title=False)

    # json_object = json.dumps(responses, indent=4)
    # with open("sample.json", "w") as outfile:
    #     outfile.write(json_object)
    f = open("sample.json")
    responses = json.load(f)
    f.close()
    return pdf_generator(responses, start)


def pdf_generator(responses, email_address, start):
    print("[!] Server logs: Result Generation started")
    document = Document()

    p = document.add_paragraph()
    p = p.insert_paragraph_before("                    ")
    r = p.add_run()
    r.add_picture("assets/report_logo.jpeg")

    ID = open("assets/ID.txt", "r").read()
    now = time.time()
    DateAndTime = now.strftime("%d/%m/%Y %H:%M:%S")

    updated_ID = int(ID) + 1
    update = open("assets/ID.txt", "w")
    update.write(str(updated_ID))
    update.close()

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = (
        "ID:"
        + str("{num:0>4}".format(num=str(ID)))
        + " Generated at : "
        + str(DateAndTime)
    )
    paragraph.style = document.styles["Header"]

    document.add_heading("Results", 0)

    document.add_heading("Gist Gmail Summary")
    document.add_paragraph("Type : " + "Gmail Summary")

    document.add_heading("Execution Time")
    summarizer_time = time.time() - start - now
    document.add_paragraph("Summarization : " + str(summarizer_time) + " seconds")

    document.add_heading("Emails summarized")
    document.add_paragraph("Number of Emails summarized : " + str(len(responses)))

    document.add_page_break()

    document.add_heading("Report", 0)

    for data in responses:
        document.add_heading("Email ID : " + data["id"], 1)
        document.add_heading("Email Subject : " + data["meta"]["Subject"], 1)
        document.add_heading("Email From : " + data["meta"]["from"], 1)
        document.add_heading("Recieved at : " + data["meta"]["Date"], 1)
        document.add_heading("Email Content", 2)
        document.add_paragraph(data["content"])

        document.add_heading("Email Summary", 2)
        document.add_paragraph(data["summary"])
        document.add_paragraph("                              ")
        document.add_paragraph(
            "======================================================================"
        )
    xl = win32com.client.Dispatch("Word.Application", pythoncom.CoInitialize())
    doc_location = "temp\Result_{num:0>4}.docx".format(num=str(ID))
    document.save(doc_location)
    convert(doc_location, output_path="temp")
    os.remove(doc_location)
    return "Result_{num:0>4}.pdf".format(num=str(ID))
